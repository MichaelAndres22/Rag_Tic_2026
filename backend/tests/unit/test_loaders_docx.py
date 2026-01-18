from pathlib import Path
import docx
from app.services.ingest.loaders import load_text

def test_load_text_docx(tmp_path: Path):
    p = tmp_path / "a.docx"
    d = docx.Document()
    d.add_paragraph("FastAPI")
    d.add_paragraph("Python")
    d.save(str(p))

    text, meta = load_text(p)
    assert "FastAPI" in text
    assert meta["type"] == "docx"